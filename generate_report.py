from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        if color:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '4472C4')
        shading.set(qn('w:color'), 'FFFFFF')
        shading.set(qn('w:val'), 'clear')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    # Data rows
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    return table

def add_status_bar(doc, label, pct, color_hex):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: {pct}%  ")
    run.bold = True
    bars = int(pct / 5)
    bar_run = p.add_run("█" * bars + "░" * (20 - bars))
    bar_run.font.color.rgb = RGBColor(int(color_hex[0:2],16), int(color_hex[2:4],16), int(color_hex[4:6],16))
    bar_run.font.size = Pt(11)

# ─── TITLE ───
doc.add_picture  # skip cover, use heading
title = doc.add_heading('HARISMA AMAZON — СТАТУС ПРОЕКТА', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x17, 0x37, 0x5E)

p = doc.add_paragraph(f'Дата отчёта: {datetime.date.today().strftime("%d.%m.%Y")}')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].italic = True

doc.add_paragraph()

# ─── 1. ОБЗОР ───
add_heading(doc, '1. Обзор проекта', 1)
doc.add_paragraph(
    'Проект направлен на запуск продаж на Amazon в четырёх европейских рынках: '
    'Германия (DE), Испания (ES), Италия (IT), Франция (FR). '
    'Три бренда: QEEP (БАДы и косметика), Harly (зоотовары), Алтея (косметика и витамины). '
    'Целевые SKU: Inositol, Magnesium Citrate, Пробиотики (человек + животные), Retinol-крем.'
)

# ─── 2. ПРОИЗВОДИТЕЛЬ ───
add_heading(doc, '2. Выбор производителя — РЕШЕНИЕ ПРИНЯТО', 1)

p = doc.add_paragraph()
run = p.add_run('✅ Выбран производитель: ERAscientifico (Латвия)')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x1F, 0x7A, 0x1F)

doc.add_paragraph(
    'После сравнительного анализа двух финалистов принято решение работать с ERAscientifico '
    'как основным производителем, МЕРИВУД (Эстония) остаётся резервным вариантом.'
)

add_table(doc,
    ['Критерий', 'ERAscientifico (Латвия) ✅', 'МЕРИВУД (Эстония)'],
    [
        ['Amazon-сертификация', '✅ Есть', '⚠️ Дороже'],
        ['Гибкость формул', '✅ Высокая', '❌ Фиксированные формулы'],
        ['Условия оплаты', '50% предоплата + отсрочка', '100% предоплата'],
        ['Посещение завода', '❌ Отказывает', '✅ Допускает'],
        ['Независимые лаб.тесты', '⚠️ Отказывает (риск!)', '✅ Разрешает'],
        ['Стоимость сертификации', '✅ Ниже', '❌ Выше'],
        ['Статус', '✅ ВЫБРАН (основной)', 'Резервный'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('⚠️  Важно: ')
run.bold = True
run.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)
p.add_run(
    'ERAscientifico отказывает в проведении независимых лабораторных тестов и посещении завода. '
    'Рекомендуется зафиксировать требования по качеству в договоре и запросить COA (Certificate of Analysis) '
    'по каждой партии перед отгрузкой.'
)

# ─── 3. СТАТУС БЛОКЕРОВ ───
add_heading(doc, '3. Критические блокеры', 1)

add_table(doc,
    ['#', 'Блокер', 'Ответственный', 'Статус'],
    [
        ['1', 'WISE-счёт для Amazon', 'Лена', '⏳ В процессе'],
        ['2', 'Amazon Seller аккаунт', 'Лена', '❌ Не открыт (ждёт банка)'],
        ['3', 'VAT Германия', 'Юрист', '⏳ Подан, ждёт подтверждения'],
        ['4', 'VAT Италия / Испания / Франция', 'Лена / AVASK', '❌ Статус неизвестен'],
        ['5', 'EUIPO — регистрация ТМ', 'Патентный поверенный', '⏳ 3 мес. ожидания возражений'],
        ['6', 'Финансовые модели — верификация', 'Команда', '⚠️ Нужна проверка'],
        ['7', 'SEO-данные (устарели, янв. 2026)', 'Виталий', '⚠️ Требует обновления'],
    ]
)

# ─── 4. ПРОГРЕСС ───
add_heading(doc, '4. Прогресс по направлениям', 1)

add_status_bar(doc, 'Юридика / корп.структура    ', 70, '4472C4')
add_status_bar(doc, 'Производство / поставщики   ', 80, '70AD47')
add_status_bar(doc, 'Продукт / анализ рынка      ', 100, '70AD47')
add_status_bar(doc, 'Финансовые модели           ', 90, 'ED7D31')
add_status_bar(doc, 'SEO / листинги              ', 40, 'ED7D31')
add_status_bar(doc, 'Amazon аккаунт              ', 0, 'FF0000')
add_status_bar(doc, 'Маркетинг / запуск          ', 0, 'FF0000')

doc.add_paragraph()

# ─── 5. ФИНАНСОВАЯ МОДЕЛЬ ───
add_heading(doc, '5. Финансовая модель (P&L — итог за год)', 1)

add_table(doc,
    ['Метрика', 'Inositol', 'Magnesium', 'Итого'],
    [
        ['Продано ед./год', '4,150', '6,000', '10,150'],
        ['Выручка нетто', '$68,706', '$118,701', '$187,407'],
        ['Расходы PPC', '$25,641', '$45,223', '$70,864'],
        ['Операц. прибыль', '-$10,690', '-$33,675', '-$44,365'],
        ['Маржа', '-15.6%', '-28.4%', '-23.7%'],
        ['TACoS (PPC/выручка)', '37.3%', '38.1%', '37.8%'],
        ['Чистая прибыль (с ФОТ)', '-$34,690', '-$57,675', '-$68,365'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Выход на безубыточность (с учётом LTV): ')
run.bold = True
p.add_run('Inositol — 6–7 мес., Magnesium — 5–6 мес. (при retention 15% / 8%)')

# ─── 6. ЮНИТ-ЭКОНОМИКА ───
add_heading(doc, '6. Юнит-экономика (по объёму упаковки)', 1)

add_table(doc,
    ['Объём (капсул)', 'Inositol COGS', 'Magnesium COGS', 'Рек. розница'],
    [
        ['30 шт', '€1.55', '€2.10', '€15.95–€20.99'],
        ['60 шт', '€1.64', '€3.10', '€19.99–€24.99'],
        ['90 шт', '€1.70', '€4.07', '€23.95–€27.99'],
        ['120 шт', '€1.80', '€5.21', '€26.99–€29.99'],
        ['180 шт', '€2.10', '€6.80', '€29.95–€35.95'],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    'Оптимальный вход: €19.99–€26.99 (конкурирует с игроками 1,000–2,000 ед/мес). '
    'Достижимая маржа: 30–40% при ценовом диапазоне €24–€30.'
)

# ─── 7. SEO ───
add_heading(doc, '7. SEO — статус по рынкам', 1)

add_table(doc,
    ['Продукт', 'Германия (DE)', 'Франция (FR)', 'Италия (IT)', 'Испания (ES)'],
    [
        ['Inositol 500mg', '✅ Готово', '✅ Готово', '✅ Готово', '❌ Нет данных'],
        ['Magnesium Citrate', '✅ Готово', '✅ Готово', '✅ Готово', '❌ Нет данных'],
        ['Пробиотики (люди)', '✅ Готово', '✅ Готово', '⚠️ Частично', '❌ Нет данных'],
        ['Пробиотики (животные)', '✅ Готово', '✅ Готово', '⚠️ Частично', '❌ Нет данных'],
        ['Retinol Cream 0.5%', '✅ Готово', '✅ Готово', '⚠️ Частично', '❌ Нет данных'],
        ['A+ Content', '⚠️ Частично', '⚠️ Частично', '⚠️ Частично', '❌ Нет данных'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Примечание: ')
run.bold = True
p.add_run('SEO-данные собраны в январе 2026 — требуют обновления перед созданием листингов.')

# ─── 8. КОНКУРЕНТНЫЙ АНАЛИЗ ───
add_heading(doc, '8. Конкурентный анализ — ключевые выводы', 1)

add_table(doc,
    ['Ценовой тир', 'Цена', 'Продажи/мес', 'Выручка/мес'],
    [
        ['Бюджет', '€11.99–€16.99', '2,000–10,000+ ед', 'до €150K'],
        ['Стандарт ✅ (наш тир)', '€19.99–€29.99', '1,000–3,000 ед', '€30K–€80K'],
        ['Премиум', '€30–€58.90', '100–500 ед', '€5K–€25K'],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    'Наиболее прибыльный диапазон для входа: €20–€35. '
    'Конкуренты с высоким объёмом (3,000–26,000 ед/мес) работают в бюджетном тире — '
    'конкурировать по цене нецелесообразно. Фокус: качество + LTV.'
)

# ─── 9. ЛОГИСТИКА ───
add_heading(doc, '9. Логистика', 1)
doc.add_paragraph(
    '1. Производство на заводе ERAscientifico (Латвия)\n'
    '2. 2 недели хранение на заводе\n'
    '3. Передача на 3PL консолидационный склад\n'
    '4. Финальная отправка на Amazon FBA (DE/FR/IT/ES)\n'
    'Маркировка: мультиязычная (FR, ES, IT, DE, EN)'
)

# ─── 10. СЛЕДУЮЩИЕ ШАГИ ───
add_heading(doc, '10. Следующие шаги (приоритет)', 1)

add_table(doc,
    ['#', 'Задача', 'Срок', 'Ответственный'],
    [
        ['1', 'Открыть WISE-счёт → разблокировать Amazon аккаунт', 'ASAP', 'Лена'],
        ['2', 'Уточнить VAT IT/ES/FR (рассмотреть AVASK)', 'ASAP', 'Лена / юрист'],
        ['3', 'Подписать договор с ERAscientifico + зафиксировать COA требования', '1–2 нед', 'Лена / технолог'],
        ['4', 'Верифицировать финансовые модели (P&L, PPC, юнит-эк.)', '1 нед', 'Команда'],
        ['5', 'Обновить SEO-данные (с янв. 2026)', '1–2 нед', 'Виталий'],
        ['6', 'Создать листинги на Amazon (после открытия аккаунта)', '2–3 нед', 'Виталий'],
        ['7', 'Запустить тестовую партию PPC (DE рынок)', 'После листингов', 'Команда'],
    ]
)

# ─── FOOTER ───
doc.add_paragraph()
p = doc.add_paragraph('Документ сформирован автоматически на основе файлов проекта.')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(9)
p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

output_path = '/home/user/Harisma-Amazon/Статус проекта Harisma Amazon.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
